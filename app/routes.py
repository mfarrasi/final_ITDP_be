from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from app.db import SessionLocal
from app.models import *
from werkzeug.security import generate_password_hash, check_password_hash
from sqlalchemy.orm.attributes import flag_modified
from sqlalchemy.orm import joinedload
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from datetime import timedelta, datetime
from docx import Document
from io import BytesIO
# from app.summarizer import summarize
import cloudinary
import cloudinary.uploader
from cloudinary.utils import cloudinary_url

# Configuration       
cloudinary.config( 
    cloud_name = "dhpmkmxfa", 
    api_key = "297779885711739", 
    api_secret = "rTz_CTSGrrxkkHEszbitCRoCVLc", # Click 'View API Keys' above to copy your API secret
    secure=True
)

app = Flask(__name__)
CORS(app)
app.config["JWT_SECRET_KEY"] = "your_secret_key"
app.config["JWT_ACCESS_TOKEN_EXPIRES"] = timedelta(days=1)
jwt = JWTManager(app)



@app.route("/users", methods=["GET", "POST"])
def users():
    db = SessionLocal()
    if request.method == "GET":
        try:
            user_list = db.query(User).all()
            result = [
        {
            "id": user.id,
            "nama": user.name,
            "email": user.email,
            "password": user.password,
            "role": user.role
        } 
        for user in user_list
            ]

            return jsonify(result)
        except Exception as e:
            return jsonify({"message": str(e)}), 500
        finally:
            db.close()

    elif request.method == "POST":
        try:
            data = request.json
            required_fields = ["name", "email", "password"]
            for field in required_fields:
                if field not in data or not data[field]:
                    return jsonify({"message": f"{field.capitalize()} tidak boleh kosong!"}), 400

            import re
            email_pattern = r"^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$"
            if not re.match(email_pattern, data["email"]):
                return jsonify({"message": "Format email tidak valid!"}), 400

            if len(data["password"]) < 8:
                return jsonify({"message": "Password harus minimal 8 karakter!"}), 400
            
            hashed_password = generate_password_hash(data["password"], method="pbkdf2:sha256")
            
            new_user = User(
                name=data["name"],
                email=data["email"],
                password=hashed_password,
                role='AO PPK'
            )
            
            db.add(new_user)
            db.commit()
            return jsonify({"message": "User berhasil ditambahkan!"})
        except Exception as e:
            return jsonify({"message": str(e)}), 500
        finally:
            db.close()

@app.route("/users/<int:user_id>", methods=["GET"])
def get_user_by_id(user_id):
    db = SessionLocal()
    try:
        user = db.query(User).filter_by(id=user_id).first()
        if not user:
            return jsonify({"message": "User tidak ditemukan!"}), 404
        
        result = {
            "id": user.id,
            "nama": user.nama,
            "email": user.email,
            "password": user.password,
            "token": create_access_token(identity=str(user.id))
        }
        return jsonify(result), 200
    except Exception as e:
        return jsonify({"message": str(e)}), 500
    finally:
        db.close()

@app.route('/login', methods=['POST'])
def login():
    print(f"Request method: {request.method}")

    db = SessionLocal()
    data = request.get_json()

    if not data:
        db.close()
        return jsonify({"error": "No input data provided"}), 400

    id = data.get('id')
    email = data.get('email')
    password = data.get('password')
    role = data.get('role')

    if not email or not password:
        db.close()
        return jsonify({"error": "Email and password are required"}), 400

    user = db.query(User).filter_by(email=email).first()
    if not user or not check_password_hash(user.password, password):
        db.close()
        return jsonify({"error": "Invalid credentials"}), 401

    access_token = create_access_token(identity=str(user.id))
    db.close()
    return jsonify({"access_token": access_token, "email": email, "role": user.role}), 200

# Update user
@app.route('/users', methods=['PUT'])
@jwt_required()
def update_user():
    db = SessionLocal()
    user_id = get_jwt_identity()
    data = request.get_json()

    user = db.query(User).filter_by(id=user_id).first()
    
    if not user:
        return jsonify({"message": "User not found"}), 404

    if 'nama' in data and data['nama']:
        user.nama = data['nama']
    if 'email' in data and data['email']:
        user.email = data['email']
    if 'password' in data and data['password']:
        user.password = generate_password_hash(data["password"])

    db.commit()
    return jsonify({"message": "User updated successfully!"}), 200

# Delete User (Requires Authentication)
@app.route('/users', methods=['DELETE'])
@jwt_required()
def delete_user():
    db = SessionLocal()
    user_id = get_jwt_identity()
    
    user = db.query(User).filter_by(id=user_id).first()
    if not user:
        return jsonify({"message": "User not found"}), 404

    db.delete(user)
    db.commit()
    db.close()
    
    return jsonify({"message": "User deleted successfully!"}), 200

    db = SessionLocal()
    try:
        user_id = get_jwt_identity()

        user = db.query(User).filter_by(id=user_id).first()
        if not user:
            return jsonify({"message": "User tidak ditemukan"}), 404

        lagu_favorite = db.query(Lagu).filter(
            Lagu.liked_by.op("@>")(f'{{{user_id}}}')
        ).all()

        result = [
            {
                "id": lagu.id,
                "judul": lagu.judul,
                "link": lagu.link,
                "gambar": lagu.gambar,
                "liked_by": lagu.liked_by,\
            }
            for lagu in lagu_favorite
        ]

        return jsonify(result), 200

    except Exception as e:
        db.rollback()
        return jsonify({"message": str(e)}), 500
    finally:
        db.close()

# Get Kantor all
@app.route('/kantor', methods=['GET'])
def get_all_kantor():
    db = SessionLocal()
    try:
        kantor_list = db.query(Kantor).all()
        result = []

        for kantor in kantor_list:
            result.append({
                "kantor_id": kantor.kantor_id,
                "kantor_cabang": kantor.cabang,
                "kantor_wilayah": kantor.wilayah,
                "jumlah_total_npl": kantor.jumlah_total_npl
            })

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# Get kantor by id
@app.route('/kantor/<int:kantor_id>', methods=['GET'])
def get_kantor_by_id(kantor_id):
    db = SessionLocal()
    try:
        kantor = db.query(Kantor).filter_by(kantor_id=kantor_id).first()
        if not kantor:
            return jsonify({"error": "Kantor not found"}), 404

        result = {
            "kantor_id": kantor.kantor_id,
            "kantor_cabang": kantor.cabang,
            "kantor_wilayah": kantor.wilayah,
            "jumlah_total_npl": kantor.jumlah_total_npl
        }

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

@app.route('/perusahaan', methods=['GET'])
@jwt_required()
def get_perusahaan():
    db = SessionLocal()

    try:
        perusahaan_list = db.query(Perusahaan).options(joinedload(Perusahaan.kantor)).all()
        result = []
        for perusahaan in perusahaan_list:
            kantor = db.query(Kantor).filter_by(kantor_id=perusahaan.kantor_id).first()
            persentase = float(perusahaan.total_outstanding/kantor.jumlah_total_npl) * 100
            persentase = round(persentase, 2)
            result.append({
                "cif": perusahaan.cif,
                "nama_perusahaan": perusahaan.nama_perusahaan,
                "total_outstanding": float(perusahaan.total_outstanding or 0),
                "total_npl": float(kantor.jumlah_total_npl) if kantor else None,
                "kantor_cabang": kantor.cabang if kantor else None,
                "kantor_wilayah": kantor.wilayah if kantor else None,
                "persentase_npl": str(persentase)+ "%"
            })

        return jsonify(result), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

@app.route('/perusahaan', methods=['POST'])
@jwt_required()
def add_perusahaan():
    db = SessionLocal()
    data = request.get_json()

    try:
        # Extract input data
        cif = data.get('cif')
        nama_perusahaan = data.get('nama_perusahaan')
        total_outstanding = data.get('total_outstanding')
        cabang = data.get('cabang')

        # Validate input
        if not all([cif, nama_perusahaan, cabang]):
            return jsonify({"error": "Missing required fields"}), 400

        # Look up kantor by cabang name
        kantor = db.query(Kantor).filter_by(cabang=cabang).first()
        if not kantor:
            return jsonify({"error": "Kantor cabang not found"}), 404

        # Create and save new Perusahaan
        new_perusahaan = Perusahaan(
            cif=cif,
            nama_perusahaan=nama_perusahaan,
            total_outstanding=total_outstanding,
            kantor_id=kantor.kantor_id
        )

        db.add(new_perusahaan)
        db.commit()

        return jsonify({"message": "Perusahaan added successfully!"}), 201

    # except IntegrityError:
    #     db.rollback()
    #     return jsonify({"error": "CIF already exists or constraint error"}), 400
    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# Update perusahaan
@app.route('/perusahaan/<cif>', methods=['PUT'])
@jwt_required()
def update_perusahaan(cif):
    db = SessionLocal()
    data = request.get_json()

    try:
        perusahaan = db.query(Perusahaan).filter_by(cif=cif).first()

        if not perusahaan:
            return jsonify({"error": "Perusahaan not found"}), 404

        # Update simple fields
        if "nama_perusahaan" in data:
            perusahaan.nama_perusahaan = data["nama_perusahaan"]

        if "total_outstanding" in data:
            perusahaan.total_outstanding = data["total_outstanding"]

        # Update kantor by name if provided
        if "kantor_cabang" in data:
            kantor = db.query(Kantor).filter_by(kantor_cabang=data["kantor_cabang"]).first()
            if not kantor:
                return jsonify({"error": "Kantor not found"}), 404
            perusahaan.kantor_id = kantor.kantor_id

        db.commit()
        return jsonify({"message": "Perusahaan updated successfully"}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# Delete perusahan
@app.route('/perusahaan/<cif>', methods=['DELETE'])
@jwt_required()
def delete_perusahaan(cif):
    db = SessionLocal()

    try:
        perusahaan = db.query(Perusahaan).filter_by(cif=cif).first()

        if not perusahaan:
            return jsonify({"error": "Perusahaan not found"}), 404

        # Optional: Delete related fasilitas first (to avoid FK errors)
        fasilitas_list = db.query(Fasilitas).filter_by(cif=cif).all()
        for fasilitas in fasilitas_list:
            # Delete agunan linked to this fasilitas
            db.query(Agunan).filter_by(deal_ref=fasilitas.deal_ref).delete()
            # Delete from user_deals if necessary
            db.query(UserDeals).filter_by(deal_ref=fasilitas.deal_ref).delete()
            # Delete event_roadmap and history
            db.query(EventRoadmap).filter_by(deal_ref=fasilitas.deal_ref).delete()
            db.query(History).filter_by(deal_ref=fasilitas.deal_ref).delete()
            db.delete(fasilitas)

        db.delete(perusahaan)
        db.commit()

        return jsonify({"message": "Perusahaan and related data deleted successfully"}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# Get fasilitas per CIF
@app.route('/perusahaan/<cif>/fasilitas', methods=['GET'])
@jwt_required()
def get_fasilitas_by_cif(cif):
    db = SessionLocal()

    try:
        perusahaan = db.query(Perusahaan).filter_by(cif=cif).first()
        if not perusahaan:
            return jsonify({"error": "Perusahaan not found"}), 404

        fasilitas_list = db.query(Fasilitas).filter_by(cif=cif).all()

        results = []
        for f in fasilitas_list:
            persentase = float(f.jumlah_outstanding/perusahaan.total_outstanding) * 100
            persentase = round(persentase, 2)
            results.append({
                "deal_ref": f.deal_ref,
                "jenis_fasilitas": f.jenis_fasilitas,
                "jumlah_outstanding": f.jumlah_outstanding,
                "tanggal_mulai_macet": f.tanggal_mulai_macet.strftime("%Y-%m-%d") if f.tanggal_mulai_macet else None,
                "progres_npl": f.progres_npl,
                "persentase": str(persentase) + "%"
            })

        return jsonify(results), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# Get semua fasilitas
@app.route('/fasilitas', methods=['GET'])
@jwt_required()
def get_all_fasilitas():
    db = SessionLocal()
    try:
        fasilitas_list = db.query(Fasilitas).all()

        result = [
            {
                "deal_ref": fasilitas.deal_ref,
                "jenis_fasilitas": fasilitas.jenis_fasilitas,
                "jumlah_outstanding": float(fasilitas.jumlah_outstanding or 0)
            }
            for fasilitas in fasilitas_list
        ]

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# Nambah Fasilitas
@app.route('/fasilitas', methods=['POST'])
@jwt_required()
def add_fasilitas():
    db = SessionLocal()
    data = request.get_json()

    try:
        required_fields = ['deal_ref', 'cif', 'ao_komersial_nama', 'ao_ppk_nama']
        if not all(data.get(field) for field in required_fields):
            return jsonify({"error": "Missing required fields"}), 400

        # Check perusahaan exists
        perusahaan = db.query(Perusahaan).filter_by(cif=data['cif']).first()
        if not perusahaan:
            return jsonify({"error": "Perusahaan (CIF) not found"}), 404

        # Look up AO Komersial
        ao_komersial = db.query(User).filter_by(name=data['ao_komersial_nama']).first()
        if not ao_komersial:
            return jsonify({"error": "AO Komersial not found"}), 404

        # Look up AO PPK
        ao_ppk = db.query(User).filter_by(name=data['ao_ppk_nama']).first()
        if not ao_ppk:
            return jsonify({"error": "AO PPK not found"}), 404

        # Create new fasilitas
        fasilitas = Fasilitas(
            deal_ref=data['deal_ref'],
            cif=data['cif'],
            jenis_fasilitas=data.get('jenis_fasilitas'),
            jumlah_outstanding=data.get('jumlah_outstanding'),
            tanggal_mulai_macet=parse_date(data.get('tanggal_mulai_macet')),
            key_person_perusahaan=data.get('key_person_perusahaan'),
            progres_npl=data.get('progres_npl'),
            restruktur_terakhir=parse_date(data.get('restruktur_terakhir')),
            ao_komersial=ao_komersial.id,
            ao_ppk=ao_ppk.id,
            jumlah_kredit_recovered=data.get('jumlah_kredit_recovered', 0)
        )

        db.add(fasilitas)
        db.commit()

        return jsonify({"message": "Fasilitas added successfully!"}), 201

    except IntegrityError:
        db.rollback()
        return jsonify({"error": "Deal ref already exists or constraint error"}), 400
    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# Helper function for date parsing
def parse_date(date_str):
    if date_str:
        try:
            return datetime.strptime(date_str, "%Y-%m-%d").date()
        except ValueError:
            raise ValueError(f"Invalid date format: {date_str}")
    return None

# Detail fasilitas
@app.route('/fasilitas/<deal_ref>', methods=['GET'])
@jwt_required()
def get_fasilitas_detail(deal_ref):
    db = SessionLocal()
    try:
        fasilitas = db.query(Fasilitas).filter_by(deal_ref=deal_ref).first()

        if not fasilitas:
            return jsonify({"error": "Fasilitas not found"}), 404

        # Related perusahaan
        perusahaan = db.query(Perusahaan).filter_by(cif=fasilitas.cif).first()

        # AOs
        ao_komersial = db.query(User).filter_by(id=fasilitas.ao_komersial).first()
        ao_ppk = db.query(User).filter_by(id=fasilitas.ao_ppk).first()

        # Optional: related agunan
        agunan_list = db.query(Agunan).filter_by(deal_ref=deal_ref).all()

        response = {
            "deal_ref": fasilitas.deal_ref,
            "cif": fasilitas.cif,
            "nama_perusahaan": perusahaan.nama_perusahaan if perusahaan else None,
            "jenis_fasilitas": fasilitas.jenis_fasilitas,
            "jumlah_outstanding": float(fasilitas.jumlah_outstanding or 0),
            "tanggal_mulai_macet": fasilitas.tanggal_mulai_macet.isoformat() if fasilitas.tanggal_mulai_macet else None,
            "key_person_perusahaan": fasilitas.key_person_perusahaan,
            "progres_npl": fasilitas.progres_npl,
            "restruktur_terakhir": fasilitas.restruktur_terakhir.isoformat() if fasilitas.restruktur_terakhir else None,
            "ao_komersial": ao_komersial.name if ao_komersial else None,
            "ao_ppk": ao_ppk.name if ao_ppk else None,
            "jumlah_kredit_recovered": float(fasilitas.jumlah_kredit_recovered or 0),
            "agunan": [
                {
                    "jenis_agunan": agunan.jenis_agunan,
                    "reappraisal_terakhir": float(agunan.reappraisal_terakhir or 0),
                    "tanggal_reappraisal": agunan.tanggal_reappraisal.isoformat() if agunan.tanggal_reappraisal else None,
                    "status_agunan": agunan.status_agunan
                }
                for agunan in agunan_list
            ]
        }

        return jsonify(response), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# Update fasilitas
@app.route('/fasilitas/<deal_ref>', methods=['PUT'])
@jwt_required()
def update_fasilitas(deal_ref):
    db = SessionLocal()
    data = request.get_json()

    try:
        fasilitas = db.query(Fasilitas).filter_by(deal_ref=deal_ref).first()
        if not fasilitas:
            return jsonify({"error": "Fasilitas not found"}), 404

        # Optional fields
        if "jenis_fasilitas" in data:
            fasilitas.jenis_fasilitas = data["jenis_fasilitas"]
        if "jumlah_outstanding" in data:
            fasilitas.jumlah_outstanding = data["jumlah_outstanding"]
        if "tanggal_mulai_macet" in data:
            fasilitas.tanggal_mulai_macet = parse_date(data["tanggal_mulai_macet"])
        if "key_person_perusahaan" in data:
            fasilitas.key_person_perusahaan = data["key_person_perusahaan"]
        if "progres_npl" in data:
            fasilitas.progres_npl = data["progres_npl"]
        if "restruktur_terakhir" in data:
            fasilitas.restruktur_terakhir = parse_date(data["restruktur_terakhir"])
        if "jumlah_kredit_recovered" in data:
            fasilitas.jumlah_kredit_recovered = data["jumlah_kredit_recovered"]

        # AO names -> lookup UID
        if "ao_komersial" in data:
            ao_k = db.query(User).filter_by(role=data["AO Komersial"]).first()
            if not ao_k:
                return jsonify({"error": "AO Komersial not found"}), 400
            fasilitas.ao_komersial = ao_k.id

        if "ao_ppk" in data:
            ao_p = db.query(User).filter_by(role=data["AO PPK"]).first()
            if not ao_p:
                return jsonify({"error": "AO PPK not found"}), 400
            fasilitas.ao_ppk = ao_p.id

        db.commit()
        return jsonify({"message": "Fasilitas updated successfully"}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# Update agunan
@app.route('/fasilitas/<deal_ref>/agunan', methods=['PUT'])
@jwt_required()
def update_agunan(deal_ref):
    db = SessionLocal()
    data = request.get_json()

    try:
        # Check that fasilitas exists
        fasilitas = db.query(Fasilitas).filter_by(deal_ref=deal_ref).first()
        if not fasilitas:
            return jsonify({"error": "Fasilitas not found"}), 404

        # Delete existing agunan entries
        db.query(Agunan).filter_by(deal_ref=deal_ref).delete()

        # Insert new agunan entries
        new_agunan_list = []
        for agunan_data in data.get("agunan_list", []):
            agunan = Agunan(
                deal_ref=deal_ref,
                jenis_agunan=agunan_data.get("jenis_agunan"),
                reappraisal_terakhir=agunan_data.get("reappraisal_terakhir"),
                tanggal_reappraisal=parse_date(agunan_data.get("tanggal_reappraisal")),
                status_agunan=agunan_data.get("status_agunan")
            )
            new_agunan_list.append(agunan)

        db.add_all(new_agunan_list)
        db.commit()

        return jsonify({"message": "Agunan updated successfully"}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# Add History
@app.route('/history', methods=['POST'])
@jwt_required()
def add_history():
    db = SessionLocal()
    user_id = get_jwt_identity()

    try:
        # Use request.form instead of request.get_json()
        deal_ref = request.form.get("deal_ref")
        jenis_kegiatan = request.form.get("jenis_kegiatan")
        keterangan_kegiatan = request.form.get("keterangan_kegiatan")
        tanggal = request.form.get("tanggal")
        status = request.form.get("status")

        # Validate required fields
        required_fields = {"deal_ref": deal_ref, "jenis_kegiatan": jenis_kegiatan}
        missing = [field for field, value in required_fields.items() if not value]
        if missing:
            return jsonify({"error": f"Missing required fields: {', '.join(missing)}"}), 400

        # Check if an image is uploaded
        image = request.files.get('image')
        image_url = None

        if image:
            result = cloudinary.uploader.upload(
                image,
                fetch_format="auto", quality="auto",
                width=500, height=500, crop="auto", gravity="auto"
            )
            image_url = result['secure_url']

        # Save to database
        new_history = History(
            deal_ref=deal_ref,
            jenis_kegiatan=jenis_kegiatan,
            ao_input=user_id,
            keterangan_kegiatan=keterangan_kegiatan,
            tanggal=tanggal,
            status=status,
            image=image_url
        )

        db.add(new_history)
        db.commit()

        return jsonify({"message": "History added successfully"}), 201

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# Get History bisa filter by AO atau deal ref cuman tambah ?ao_input= / ?deal_ref=
@app.route('/history', methods=['GET'])
@jwt_required()
def get_history():
    db = SessionLocal()
    try:
        # Optional query params
        deal_ref = request.args.get("deal_ref")
        ao_name = request.args.get("ao_input")

        query = db.query(History, User.name).join(User, History.ao_input == User.id)

        if deal_ref:
            query = query.filter(History.deal_ref == deal_ref)

        if ao_name:
            query = query.filter(User.name.ilike(f"%{ao_name}%"))

        results = query.all()

        history_list = []
        for history, ao_nama in results:
            history_list.append({
                "event_history_id": history.event_history_id,
                "deal_ref": history.deal_ref,
                "jenis_kegiatan": history.jenis_kegiatan,
                "ao_input": ao_nama,
                "keterangan_kegiatan": history.keterangan_kegiatan,
                "tanggal": history.tanggal.isoformat() if history.tanggal else None,
                "status": history.status,
                "image": history.image
            })

        return jsonify(history_list), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

#get history by deal ref
@app.route('/history/<deal_ref>', methods=['GET'])
@jwt_required()
def get_history_by_id(deal_ref):
    db = SessionLocal()
    try:
        histories = db.query(History).filter_by(deal_ref=deal_ref).all()
        if not histories:
            return jsonify({"message": "No history records found for this deal_ref"}), 404

        result = []
        for history in histories:
            ao = db.query(User).filter_by(id=history.ao_input).first()
            result.append({
                "event_history_id": history.event_history_id,
                "deal_ref": history.deal_ref,
                "jenis_kegiatan": history.jenis_kegiatan,
                "ao_input": ao.name if ao else None,
                "keterangan_kegiatan": history.keterangan_kegiatan,
                "tanggal": history.tanggal,
                "status": history.status
            })

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# update history by history id
@app.route('/history/<int:event_history_id>', methods=['PUT'])
@jwt_required()
def update_history(event_history_id):
    db = SessionLocal()
    data = request.get_json()
    user_id = get_jwt_identity()

    try:
        history = db.query(History).filter_by(event_history_id=event_history_id).first()
        if not history:
            return jsonify({"error": "History record not found"}), 404

        if 'deal_ref' in data:
            history.deal_ref = data['deal_ref']
        if 'jenis_kegiatan' in data:
            history.jenis_kegiatan = data['jenis_kegiatan']
        if 'keterangan_kegiatan' in data:
            history.keterangan_kegiatan = data['keterangan_kegiatan']
        if 'tanggal' in data:
            history.tanggal = data['tanggal']
        if 'ao_input' in data:
            history.ao_input = user_id
        if 'status' in data:
            history.status = data['status']

        db.commit()
        return jsonify({"message": "History updated successfully"}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# delete history by history id
@app.route('/history/<int:event_history_id>', methods=['DELETE'])
@jwt_required()
def delete_history(event_history_id):
    db = SessionLocal()
    try:
        history = db.query(History).filter_by(event_history_id=event_history_id).first()
        if not history:
            return jsonify({"error": "History record not found"}), 404

        db.delete(history)
        db.commit()
        return jsonify({"message": "History record deleted successfully"}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500

    finally:
        db.close()

# get all roadmap plan
@app.route("/roadmap-plans", methods=["GET"])
@jwt_required()
def get_all_roadmap_plans():
    db = SessionLocal()
    try:
        plans = db.query(RoadmapPlan)

        result = []
        for plan in plans:
            ao = db.query(User).filter_by(id=plan.created_by).first()
            result.append({
                "plan_id": plan.plan_id,
                "deal_ref": plan.deal_ref,
                "status": plan.status,
                "created_at": plan.created_at,
                "updated_at": plan.updated_at,
                "created_by": {
                    "uid": ao.id,
                    "name": ao.name,
                    "email": ao.email
                },
                "event_count": len(plan.events)
            })

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# get roadmap plan by plan id
@app.route("/roadmap-plans/<int:plan_id>", methods=["GET"])
@jwt_required()
def get_events_by_plan(plan_id):
    db = SessionLocal()
    try:
        plan = db.query(RoadmapPlan).filter_by(plan_id=plan_id).first()
        if not plan:
            return jsonify({"message": "Roadmap plan not found"}), 404

        events = db.query(RoadmapEvent).filter_by(plan_id=plan_id).all()

        result = []
        for event in events:
            result.append({
                "event_id": event.event_id,
                "jenis_kegiatan": event.jenis_kegiatan,
                "ao_input": event.ao_input,
                "cif": event.cif,
                "keterangan_kegiatan": event.keterangan_kegiatan,
                "tanggal": event.tanggal,
                "update_terakhir": event.update_terakhir
            })

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# add roadmap events nanti jadi 1 plan
@app.route("/roadmap-plans", methods=["POST"])
@jwt_required()
def create_roadmap_plan_with_events():
    db = SessionLocal()
    data = request.get_json()
    user_id = get_jwt_identity()

    try:
        # Required fields
        deal_ref = data.get("deal_ref")
        cif = data.get('cif')
        created_by = user_id
        events = data.get("events", [])  # List of roadmap events

        if not deal_ref or not isinstance(events, list):
            return jsonify({"message": "Missing required fields"}), 400

        # Create roadmap plan
        plan = RoadmapPlan(
            deal_ref=deal_ref,
            created_by=created_by,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        db.add(plan)
        db.commit()  # Commit to get plan_id
        db.refresh(plan)

        # Create roadmap events
        event_objects = []
        for item in events:
            event = RoadmapEvent(
                plan_id=plan.plan_id,
                jenis_kegiatan=item.get("jenis_kegiatan"),
                ao_input=user_id,
                cif=cif,
                keterangan_kegiatan=item.get("keterangan_kegiatan"),
                tanggal=datetime.strptime(item.get("tanggal"), "%Y-%m-%d") if item.get("tanggal") else None
            )
            event_objects.append(event)

        db.add_all(event_objects)
        db.commit()

        return jsonify({
            "message": "Roadmap plan and events created successfully.",
            "plan_id": plan.plan_id,
            "event_count": len(event_objects)
        }), 201

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# Mungkin perlu di fix
# update roadmap plan
@app.route("/roadmap-plans/<int:plan_id>", methods=["PUT"])
@jwt_required()
def update_roadmap_plan_and_events(plan_id):
    db = SessionLocal()
    data = request.get_json()

    try:
        plan = db.query(RoadmapPlan).filter_by(plan_id=plan_id).first()
        if not plan:
            return jsonify({"message": "Roadmap plan not found"}), 404

        # --- Update Plan ---
        if "deal_ref" in data:
            plan.deal_ref = data["deal_ref"]
        if "status" in data:
            plan.status = data["status"]
        plan.updated_at = datetime.utcnow()

        # --- Update Events ---
        if "events" in data:
            for event_data in data["events"]:
                event_id = event_data.get("event_id")
                if not event_id:
                    continue  # skip if no ID provided

                event = db.query(RoadmapEvent).filter_by(event_id=event_id, plan_id=plan_id).first()
                if not event:
                    continue  # skip if not found

                # Update fields if present
                if "jenis_kegiatan" in event_data:
                    event.jenis_kegiatan = event_data["jenis_kegiatan"]
                if "cif" in event_data:
                    event.cif = event_data["cif"]
                if "keterangan_kegiatan" in event_data:
                    event.keterangan_kegiatan = event_data["keterangan_kegiatan"]
                if "tanggal" in event_data:
                    event.tanggal = datetime.strptime(event_data["tanggal"], "%Y-%m-%d")
                event.update_terakhir = datetime.utcnow()

        db.commit()
        return jsonify({"message": "Roadmap plan and events updated successfully."}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# delete roadmap PLAN by plan id
@app.route("/roadmap-plans/<int:plan_id>", methods=["DELETE"])
@jwt_required()
def delete_roadmap_plan(plan_id):
    db = SessionLocal()

    try:
        plan = db.query(RoadmapPlan).filter_by(plan_id=plan_id).first()

        if not plan:
            return jsonify({"message": "Roadmap plan not found"}), 404

        # Optional: delete associated events first
        db.query(RoadmapEvent).filter_by(plan_id=plan_id).delete()

        # Then delete the plan
        db.delete(plan)
        db.commit()

        return jsonify({"message": "Roadmap plan and associated events deleted successfully."}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# delete roadmap EVENT by event id
@app.route("/roadmap-events/<int:event_id>", methods=["DELETE"])
@jwt_required()
def delete_roadmap_event(event_id):
    db = SessionLocal()

    try:
        event = db.query(RoadmapEvent).filter_by(event_id=event_id).first()

        if not event:
            return jsonify({"message": "Roadmap event not found"}), 404

        db.delete(event)
        db.commit()

        return jsonify({"message": "Roadmap event deleted successfully."}), 200

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

# get semua roadmap event by deal ref yang status Accepted
@app.route("/roadmap-events/<string:deal_ref>", methods=["GET"])
def get_events_of_accepted_roadmap_plans(deal_ref):
    db = SessionLocal()
    try:
        events = (
            db.query(RoadmapEvent)
            .join(RoadmapPlan)
            .filter(RoadmapPlan.deal_ref == deal_ref, RoadmapPlan.status == "Accepted", RoadmapEvent.cif == RoadmapPlan.cif)
            .all()
        )

        result = []
        for event in events:
            result.append({
                "event_id": event.event_id,
                "jenis_kegiatan": event.jenis_kegiatan,
                "tanggal": event.tanggal.strftime('%Y-%m-%d') if event.tanggal else None,
                "keterangan": event.keterangan_kegiatan,
                "cif": event.cif,
                "plan": event.plan.status,
                "ao_input": event.ao.name
            })

        return jsonify(result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()

#perlu fix
# geenerate docs
@app.route("/perusahaan/<string:cif>/docx", methods=["GET"])
def generate_docx(cif):
    db = SessionLocal()
    perusahaan = db.query(Perusahaan).filter_by(cif=cif).first()

    if not perusahaan:
        return {"message": "Perusahaan not found"}, 404

    kantor = db.query(Kantor).filter_by(kantor_id=perusahaan.kantor_id).first()
    fasilitas_list = db.query(Fasilitas).filter_by(cif=cif).all()

    # Create Word document
    doc = Document()
    doc.add_heading('Perusahaan Report', 0)

    doc.add_paragraph(f"CIF: {perusahaan.cif}")
    doc.add_paragraph(f"Nama Perusahaan: {perusahaan.nama_perusahaan}")
    doc.add_paragraph(f"Total Outstanding: {perusahaan.total_outstanding}")
    doc.add_paragraph(f"Kantor Cabang: {kantor.cabang}")
    doc.add_paragraph(f"Kantor Wilayah: {kantor.wilayah}")

    doc.add_heading('Fasilitas', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Deal Ref'
    hdr_cells[1].text = 'Jenis Fasilitas'
    hdr_cells[2].text = 'Jumlah Outstanding'

    for f in fasilitas_list:
        row_cells = table.add_row().cells
        row_cells[0].text = f.deal_ref
        row_cells[1].text = f.jenis_fasilitas or ""
        row_cells[2].text = str(f.jumlah_outstanding or "")

    # Save to memory buffer
    f = BytesIO()
    doc.save(f)
    f.seek(0)

    db.close()

    return send_file(
        f,
        as_attachment=True,
        download_name=f"perusahaan_{cif}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# upload image
# redundan
# @app.route('/upload-image', methods=['POST'])
# def upload_image():
#     db = SessionLocal()

#     try:
#         # Upload an image
#         if 'image' not in request.files:
#             return jsonify({'error': 'No image uploaded'}), 400

#         image = request.files['image']
#         result = cloudinary.uploader.upload(
#             image,
#             fetch_format="auto", quality="auto",
#             width=500, height=500, crop="auto", gravity="auto"
#             )
        
#         db.add(new_history)
#         db.commit()

#         return jsonify({'url': result['secure_url']}), 200
#     except Exception as e:
#         db.rollback()
#         return jsonify({"error": str(e)}), 500
#     finally:
#         db.close()


# @app.route("/summarize", methods=['GET'])
# def summarize_text():
#     return summarize()